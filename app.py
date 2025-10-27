import streamlit as st
import google.generativeai as genai
from pathlib import Path
import PyPDF2
import docx
import io
import json
import time
from datetime import datetime
import hashlib
import os
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# --- Page Configuration ---
st.set_page_config(
    page_title="Advanced Gemini Document Analyzer",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- Custom CSS for Professional Look ---
st.markdown("""
    <style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 0.5rem;
    }
    .stButton>button {
        width: 100%;
    }
    .chat-message {
        padding: 1rem;
        border-radius: 0.5rem;
        margin-bottom: 1rem;
    }
    .success-badge {
        background-color: #10b981;
        color: white;
        padding: 0.25rem 0.75rem;
        border-radius: 1rem;
        font-size: 0.875rem;
        display: inline-block;
    }
    </style>
""", unsafe_allow_html=True)

# --- Enhanced Helper Functions ---

def calculate_file_hash(file_content):
    """Calculate MD5 hash of file content for caching."""
    return hashlib.md5(file_content).hexdigest()

def extract_text_from_pdf(file):
    """Enhanced PDF extraction with better error handling and metadata."""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        metadata = {
            "pages": len(pdf_reader.pages),
            "author": pdf_reader.metadata.get('/Author', 'Unknown') if pdf_reader.metadata else 'Unknown',
            "title": pdf_reader.metadata.get('/Title', 'Unknown') if pdf_reader.metadata else 'Unknown'
        }
        
        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text:
                text += f"\n--- Page {page_num + 1} ---\n{page_text}"
        
        return text, metadata
    except Exception as e:
        st.error(f"PDF Error: {str(e)}")
        return None, None

def extract_text_from_docx(file):
    """Enhanced DOCX extraction with structure preservation."""
    try:
        doc = docx.Document(file)
        text = ""
        metadata = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables)
        }
        
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract tables
        if doc.tables:
            text += "\n--- Tables ---\n"
            for table_num, table in enumerate(doc.tables):
                text += f"\nTable {table_num + 1}:\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text += row_text + "\n"
        
        return text, metadata
    except Exception as e:
        st.error(f"DOCX Error: {str(e)}")
        return None, None

def extract_text_from_txt(file):
    """Enhanced TXT extraction with encoding detection."""
    try:
        content = file.read()
        # Try UTF-8 first, fallback to latin-1
        try:
            text = content.decode('utf-8')
        except UnicodeDecodeError:
            text = content.decode('latin-1')
        
        metadata = {
            "lines": len(text.split('\n')),
            "words": len(text.split())
        }
        return text, metadata
    except Exception as e:
        st.error(f"TXT Error: {str(e)}")
        return None, None

def extract_text(file):
    """Main extraction router with metadata."""
    file_extension = Path(file.name).suffix.lower()
    
    extractors = {
        '.pdf': extract_text_from_pdf,
        '.docx': extract_text_from_docx,
        '.txt': extract_text_from_txt
    }
    
    if file_extension in extractors:
        return extractors[file_extension](file)
    else:
        st.error(f"Unsupported format: {file_extension}")
        return None, None

def chunk_text(text, chunk_size=10000, overlap=500):
    """Split text into overlapping chunks for better context."""
    chunks = []
    start = 0
    text_len = len(text)
    
    while start < text_len:
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        start = end - overlap
    
    return chunks

def get_gemini_response(prompt, api_key, model_name, temperature=0.7, max_retries=3):
    """Enhanced Gemini API call with retry logic and better error handling."""
    for attempt in range(max_retries):
        try:
            genai.configure(api_key=api_key)
            
            generation_config = {
                "temperature": temperature,
                "top_p": 0.95,
                "top_k": 40,
                "max_output_tokens": 8192,
            }

            safety_settings = [
                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_ONLY_HIGH"},
                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_ONLY_HIGH"},
                {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_ONLY_HIGH"},
                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_ONLY_HIGH"},
            ]
            
            model = genai.GenerativeModel(
                model_name=model_name,
                generation_config=generation_config,
                safety_settings=safety_settings
            )
            
            response = model.generate_content(prompt)
            return response.text
            
        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(2 ** attempt)  # Exponential backoff
                continue
            else:
                error_msg = str(e)
                if "API_KEY" in error_msg.upper():
                    return "❌ Invalid API key. Please check your Google AI API key."
                elif "QUOTA" in error_msg.upper():
                    return "❌ API quota exceeded. Please try again later."
                elif "404" in error_msg:
                    return f"❌ Model '{model_name}' not found. Try 'gemini-1.5-flash' or 'gemini-1.5-pro'."
                else:
                    return f"❌ Error: {error_msg}"

def create_chat_session(api_key, model_name, document_context):
    """Initialize a persistent chat session with document context."""
    try:
        genai.configure(api_key=api_key)
        
        model = genai.GenerativeModel(
            model_name=model_name,
            generation_config={
                "temperature": 0.6,
                "top_p": 0.95,
                "max_output_tokens": 4096,
            },
            safety_settings=[
                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_ONLY_HIGH"},
                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_ONLY_HIGH"},
                {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_ONLY_HIGH"},
                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_ONLY_HIGH"},
            ],
        )
        
        # Initialize with document context
        system_instruction = f"""You are an intelligent document analysis assistant. You have access to the following document:

DOCUMENT CONTENT:
{document_context[:30000]}

Your responsibilities:
1. Answer questions ONLY based on the document content
2. Provide specific quotes and page references when available
3. If information is not in the document, clearly state: "This information is not available in the provided document."
4. Be concise but thorough in your responses
5. Use markdown formatting for better readability"""

        chat = model.start_chat(history=[])
        # Send system instruction as first message
        chat.send_message(system_instruction)
        
        return chat
    except Exception as e:
        st.error(f"Chat initialization error: {str(e)}")
        return None

# --- Session State Initialization ---
def init_session_state():
    """Initialize all session state variables."""
    defaults = {
        'extracted_text': None,
        'file_metadata': None,
        'analysis_result': None,
        'chat_session': None,
        'chat_messages': [],
        'processing_history': [],
        'current_file_hash': None,
        'analysis_cache': {}
    }
    
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

init_session_state()

# --- Sidebar Configuration ---
with st.sidebar:
    st.markdown("## ⚙️ Configuration")
    
    # Try to load API key from environment variable or config file
    default_api_key = os.getenv('GOOGLE_API_KEY', '')
    
    # If no env variable, try loading from config.json
    if not default_api_key:
        try:
            if os.path.exists('config.json'):
                with open('config.json', 'r') as f:
                    config = json.load(f)
                    default_api_key = config.get('GOOGLE_API_KEY', '')
        except:
            pass
    
    # API Key input with saved value
    if default_api_key:
        st.success("✅ API Key Loaded Successfully!")
        api_key = default_api_key
        show_key_input = st.checkbox("🔑 Change API Key", value=False)
        
        if show_key_input:
            new_api_key = st.text_input(
                "Enter New API Key",
                type="password",
                help="Get your key from https://aistudio.google.com/app/apikey"
            )
            
            if st.button("💾 Save New Key"):
                try:
                    with open('config.json', 'w') as f:
                        json.dump({'GOOGLE_API_KEY': new_api_key}, f)
                    st.success("✅ API Key saved! Refresh the page.")
                    st.rerun()
                except Exception as e:
                    st.error(f"Error saving key: {e}")
    else:
        st.warning("⚠️ No API Key Found - Please enter your key")
        api_key = st.text_input(
            "🔑 Google AI API Key",
            type="password",
            help="Get your key from https://aistudio.google.com/app/apikey"
        )
        
        if api_key:
            save_key = st.checkbox("💾 Save API Key for future use")
            
            if save_key and st.button("Save Key"):
                try:
                    # Save to config.json
                    with open('config.json', 'w') as f:
                        json.dump({'GOOGLE_API_KEY': api_key}, f)
                    st.success("✅ API Key saved successfully! Refresh the page.")
                    time.sleep(1)
                    st.rerun()
                except Exception as e:
                    st.error(f"Error saving key: {e}")
    
    st.markdown("---")
    
    # App Mode Selection
    app_mode = st.radio(
        "📱 Application Mode",
        ["🔍 Analyze Document", "💬 Chat with Document", "📊 Advanced Analytics"],
        help="Choose your interaction mode"
    )
    
    # Model Selection with descriptions
    st.markdown("### 🤖 Model Selection")
    model_options = {
        "models/gemini-2.0-flash": "Latest 2.0 Flash - Fast & efficient (recommended)",
        "models/gemini-1.5-flash": "Gemini 1.5 Flash - Reliable",
        "models/gemini-1.5-pro": "Gemini 1.5 Pro - Most capable"
    }
    
    selected_model = st.selectbox(
        "Choose model:",
        list(model_options.keys()),
        format_func=lambda x: model_options[x]
    )
    
    # Temperature slider for creativity control
    temperature = st.slider(
        "🌡️ Temperature (creativity)",
        min_value=0.0,
        max_value=1.0,
        value=0.7,
        step=0.1,
        help="Higher = more creative, Lower = more focused"
    )
    
    st.markdown("---")
    
    # Analysis Options for Analyze mode
    if "Analyze" in app_mode:
        st.markdown("### 📋 Analysis Options")
        
        analysis_type = st.selectbox(
            "Analysis type:",
            [
                "📝 Comprehensive Summary",
                "🎯 Key Points Extraction",
                "😊 Sentiment Analysis",
                "❓ Q&A Generation",
                "🏷️ Entity Recognition",
                "📑 Document Classification",
                "🔗 Topic Modeling",
                "✍️ Custom Prompt"
            ]
        )
        
        if "Custom" in analysis_type:
            custom_prompt = st.text_area(
                "Enter your prompt:",
                placeholder="e.g., Extract all dates, names, and locations...",
                height=100
            )
    
    st.markdown("---")
    
    # Document Statistics
    if st.session_state.extracted_text:
        st.markdown("### 📊 Document Stats")
        text = st.session_state.extracted_text
        st.metric("Characters", f"{len(text):,}")
        st.metric("Words", f"{len(text.split()):,}")
        st.metric("Lines", f"{len(text.split(chr(10))):,}")
        
        if st.session_state.file_metadata:
            st.json(st.session_state.file_metadata)
    
    st.markdown("---")
    st.markdown("### 📄 Supported Formats")
    st.markdown("• PDF (.pdf)\n• Word (.docx)\n• Text (.txt)")
    
    st.markdown("---")
    st.info("💡 **Tip:** Use 'Chat' mode for interactive Q&A!")

# --- Main Header ---
st.markdown('<h1 class="main-header">🤖 Advanced Gemini Document Analyzer</h1>', unsafe_allow_html=True)
st.markdown("### Upload, analyze, and chat with your documents using Google's Gemini AI")

# --- Main Content Area ---
col1, col2 = st.columns([1, 1.5])

# --- Upload Section ---
with col1:
    st.markdown("## 📤 Document Upload")
    
    uploaded_file = st.file_uploader(
        "Choose your document",
        type=['pdf', 'docx', 'txt'],
        help="Max size: 200MB per file",
        key="file_uploader"
    )
    
    if uploaded_file:
        # Display file info
        file_size = len(uploaded_file.getvalue()) / 1024 / 1024
        col_a, col_b = st.columns(2)
        col_a.metric("File Name", uploaded_file.name)
        col_b.metric("Size", f"{file_size:.2f} MB")
        
        # Process button
        process_btn = st.button("🚀 Process Document", type="primary", use_container_width=True)
        
        if process_btn:
            if not api_key:
                st.error("⚠️ Please enter your Google AI API key first!")
            else:
                # Calculate file hash for caching
                file_content = uploaded_file.getvalue()
                file_hash = calculate_file_hash(file_content)
                
                # Check if already processed
                if file_hash == st.session_state.current_file_hash:
                    st.info("ℹ️ This document is already processed!")
                else:
                    with st.spinner("🔄 Processing document..."):
                        progress_bar = st.progress(0)
                        
                        # Extract text
                        progress_bar.progress(30)
                        uploaded_file.seek(0)
                        extracted_text, metadata = extract_text(uploaded_file)
                        
                        if extracted_text:
                            progress_bar.progress(70)
                            
                            # Update session state
                            st.session_state.extracted_text = extracted_text
                            st.session_state.file_metadata = metadata
                            st.session_state.current_file_hash = file_hash
                            st.session_state.analysis_result = None
                            st.session_state.chat_session = None
                            st.session_state.chat_messages = []
                            
                            # Add to history
                            st.session_state.processing_history.append({
                                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                "filename": uploaded_file.name,
                                "size": f"{file_size:.2f} MB",
                                "chars": len(extracted_text)
                            })
                            
                            progress_bar.progress(100)
                            st.success(f"✅ Successfully processed '{uploaded_file.name}'!")
                            st.balloons()
                        else:
                            st.error("❌ Failed to extract text from document")
                            progress_bar.empty()
        
        # Show extracted text preview
        if st.session_state.extracted_text:
            with st.expander("👁️ View Document Preview", expanded=False):
                preview_length = st.slider("Preview length (chars)", 500, 5000, 1000, 500)
                st.text_area(
                    "Document content:",
                    st.session_state.extracted_text[:preview_length] + "\n\n... [truncated]",
                    height=300,
                    disabled=True
                )
                
                # Download full text
                st.download_button(
                    "⬇️ Download Full Text",
                    st.session_state.extracted_text,
                    file_name=f"{uploaded_file.name}_extracted.txt",
                    mime="text/plain",
                    use_container_width=True
                )

# --- Results Section ---
with col2:
    if not st.session_state.extracted_text:
        st.info("👈 Upload and process a document to get started")
        
        # Show processing history
        if st.session_state.processing_history:
            st.markdown("### 📜 Recent Activity")
            for item in reversed(st.session_state.processing_history[-5:]):
                st.markdown(f"**{item['filename']}** - {item['timestamp']}")
    else:
        # Create tabs based on mode
        if "Analyze" in app_mode:
            tab1, tab2 = st.tabs(["📊 Analysis Results", "💬 Quick Chat"])
            
            with tab1:
                st.markdown("## Analysis Results")
                
                if st.button("▶️ Run Analysis", type="primary", use_container_width=True):
                    if not api_key:
                        st.error("⚠️ API key required!")
                    else:
                        with st.spinner(f"🤖 Analyzing with {selected_model}..."):
                            # Prepare prompt based on analysis type
                            prompts = {
                                "📝 Comprehensive Summary": """Provide a comprehensive summary of this document including:
                                1. Main topics and themes
                                2. Key arguments or findings
                                3. Important details and data points
                                4. Conclusions or recommendations""",
                                
                                "🎯 Key Points Extraction": """Extract and organize the key points from this document:
                                - List 8-12 most important points
                                - Use clear, concise language
                                - Include specific details and data
                                - Organize by importance or theme""",
                                
                                "😊 Sentiment Analysis": """Analyze the sentiment and tone of this document:
                                - Overall sentiment (positive/negative/neutral)
                                - Emotional tone and language style
                                - Key emotional moments or shifts
                                - Target audience and intent""",
                                
                                "❓ Q&A Generation": """Generate 10 important questions and detailed answers based on this document.
                                Format as:
                                Q1: [Question]
                                A1: [Detailed answer with context]""",
                                
                                "🏷️ Entity Recognition": """Extract and categorize all entities:
                                - People: names and roles
                                - Organizations: companies, institutions
                                - Locations: places, addresses
                                - Dates: important dates and timeframes
                                - Numbers: significant statistics""",
                                
                                "📑 Document Classification": """Classify and describe this document:
                                - Document type and purpose
                                - Target audience
                                - Formality level
                                - Subject area/domain
                                - Key characteristics""",
                                
                                "🔗 Topic Modeling": """Identify and analyze the main topics:
                                - List 5-7 major topics
                                - Explain each topic with examples
                                - Show relationships between topics
                                - Identify subtopics and themes"""
                            }
                            
                            if "Custom" in analysis_type:
                                prompt_text = custom_prompt
                            else:
                                prompt_text = prompts.get(analysis_type, "Analyze this document:")
                            
                            # Chunk text if too long
                            if len(st.session_state.extracted_text) > 25000:
                                chunks = chunk_text(st.session_state.extracted_text, 20000, 1000)
                                st.info(f"📦 Processing document in {len(chunks)} chunks...")
                                
                                results = []
                                for i, chunk in enumerate(chunks):
                                    st.write(f"Processing chunk {i+1}/{len(chunks)}...")
                                    full_prompt = f"{prompt_text}\n\n--- DOCUMENT (Part {i+1}) ---\n{chunk}"
                                    result = get_gemini_response(full_prompt, api_key, selected_model, temperature)
                                    if result and not result.startswith("❌"):
                                        results.append(result)
                                
                                if results:
                                    # Combine results
                                    combined = "\n\n".join(results)
                                    # Final synthesis
                                    synthesis_prompt = f"Synthesize these analyses into one coherent response:\n\n{combined}"
                                    final_result = get_gemini_response(synthesis_prompt, api_key, selected_model, temperature)
                                    st.session_state.analysis_result = final_result
                            else:
                                full_prompt = f"{prompt_text}\n\n--- DOCUMENT ---\n{st.session_state.extracted_text}"
                                result = get_gemini_response(full_prompt, api_key, selected_model, temperature)
                                st.session_state.analysis_result = result
                
                # Display results
                if st.session_state.analysis_result:
                    st.markdown("### 🎯 Analysis Output")
                    
                    if st.session_state.analysis_result.startswith("❌"):
                        st.error(st.session_state.analysis_result)
                    else:
                        st.markdown(st.session_state.analysis_result)
                        
                        # Action buttons
                        col_a, col_b = st.columns(2)
                        with col_a:
                            st.download_button(
                                "💾 Save Analysis",
                                st.session_state.analysis_result,
                                file_name=f"analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                                mime="text/markdown",
                                use_container_width=True
                            )
                        with col_b:
                            if st.button("🔄 New Analysis", use_container_width=True):
                                st.session_state.analysis_result = None
                                st.rerun()
            
            with tab2:
                st.markdown("## 💬 Quick Questions")
                st.markdown("*Ask quick questions about your document*")
                
                quick_question = st.text_input("Your question:", key="quick_q")
                
                if st.button("Ask", use_container_width=True) and quick_question:
                    with st.spinner("🤔 Thinking..."):
                        prompt = f"Based on this document, answer this question concisely:\n\nQuestion: {quick_question}\n\nDocument:\n{st.session_state.extracted_text[:15000]}"
                        answer = get_gemini_response(prompt, api_key, selected_model, 0.5)
                        st.markdown("**Answer:**")
                        st.info(answer)
        
        elif "Chat" in app_mode:
            st.markdown("## 💬 Document Chat Interface")
            
            # Initialize chat if needed
            if st.session_state.chat_session is None and api_key:
                with st.spinner("🔧 Initializing chat..."):
                    st.session_state.chat_session = create_chat_session(
                        api_key,
                        selected_model,
                        st.session_state.extracted_text
                    )
                    if st.session_state.chat_session:
                        st.success("✅ Chat ready!")
            
            # Display chat messages
            for msg in st.session_state.chat_messages:
                with st.chat_message(msg["role"]):
                    st.markdown(msg["content"])
            
            # Chat input
            if user_input := st.chat_input("Ask anything about your document..."):
                if not api_key:
                    st.error("⚠️ API key required")
                elif not st.session_state.chat_session:
                    st.error("⚠️ Chat session not initialized")
                else:
                    # Add user message
                    st.session_state.chat_messages.append({"role": "user", "content": user_input})
                    with st.chat_message("user"):
                        st.markdown(user_input)
                    
                    # Get response
                    with st.chat_message("assistant"):
                        with st.spinner("💭 Thinking..."):
                            try:
                                response = st.session_state.chat_session.send_message(user_input)
                                answer = response.text
                                st.markdown(answer)
                                st.session_state.chat_messages.append({"role": "assistant", "content": answer})
                            except Exception as e:
                                error_msg = f"❌ Error: {str(e)}"
                                st.error(error_msg)
            
            # Chat controls
            col_a, col_b = st.columns(2)
            with col_a:
                if st.button("🔄 Reset Chat", use_container_width=True):
                    st.session_state.chat_session = None
                    st.session_state.chat_messages = []
                    st.rerun()
            with col_b:
                if st.button("💾 Export Chat", use_container_width=True):
                    chat_export = "\n\n".join([f"{m['role'].upper()}: {m['content']}" for m in st.session_state.chat_messages])
                    st.download_button(
                        "⬇️ Download",
                        chat_export,
                        file_name=f"chat_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
        
        elif "Advanced" in app_mode:
            st.markdown("## 📊 Advanced Analytics Dashboard")
            
            if st.button("🔬 Run Advanced Analysis", type="primary", use_container_width=True):
                with st.spinner("🧪 Running advanced analytics..."):
                    # Multi-analysis
                    analyses = ["Summary", "Key Points", "Entities", "Sentiment"]
                    results = {}
                    
                    progress = st.progress(0)
                    for i, analysis in enumerate(analyses):
                        prompt_map = {
                            "Summary": "Provide a 3-paragraph summary",
                            "Key Points": "List 5 key points in bullets",
                            "Entities": "Extract all named entities (people, orgs, locations)",
                            "Sentiment": "Analyze sentiment and tone in 2-3 sentences"
                        }
                        
                        prompt = f"{prompt_map[analysis]}\n\nDocument:\n{st.session_state.extracted_text[:20000]}"
                        result = get_gemini_response(prompt, api_key, selected_model, temperature)
                        results[analysis] = result
                        progress.progress((i + 1) / len(analyses))
                    
                    # Display results in tabs
                    tabs = st.tabs(analyses)
                    for tab, analysis in zip(tabs, analyses):
                        with tab:
                            st.markdown(results[analysis])

# --- Footer ---
st.markdown("---")
st.markdown("""
    <div style='text-align: center; padding: 20px;'>
        <p style='color: #666;'>
            🚀 <strong>Advanced Gemini Document Analyzer</strong> | 
            Built with Streamlit & Google Gemini | 
            <a href='https://aistudio.google.com/app/apikey' target='_blank'>Get API Key</a>
        </p>
        <p style='color: #888; font-size: 0.9rem;'>
            💡 Tip: For best results, use clear questions and select the appropriate analysis type
        </p>
    </div>
""", unsafe_allow_html=True)